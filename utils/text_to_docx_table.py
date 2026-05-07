"""
Convert summary_report.txt files to formatted DOCX tables for thesis.

Tables are extracted in the order A1–A22 as specified in the thesis appendix.
Literature baseline rows (Mode contains "Literature") have their model name underlined.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Table specification: defines A1–A22 in order.
# Each entry maps a thesis table label to a source report and table number.
# table_num == 1  → first data table (no "TABLE 1:" marker in .txt).
# table_num >= 2  → located by "TABLE {n}:" marker in .txt.
# ---------------------------------------------------------------------------
TABLE_SPEC = [
    {"label": "A1",  "title": "N = 1 Model Benchmarking Performance Metrics",       "report": "initial_1",     "table_num": 1},
    {"label": "A2",  "title": "N = 1 Model Benchmarking Reliability Metrics",        "report": "initial_1",     "table_num": 2},
    {"label": "A3",  "title": "N = 1 Model Benchmarking Efficiency Metrics",         "report": "initial_1",     "table_num": 3},
    {"label": "A4",  "title": "N = 3 Model Benchmarking Performance Metrics",        "report": "initial_3",     "table_num": 1},
    {"label": "A5",  "title": "N = 3 Model Benchmarking Reliability Metrics",        "report": "initial_3",     "table_num": 2},
    {"label": "A6",  "title": "N = 3 Model Benchmarking Efficiency Metrics",         "report": "initial_3",     "table_num": 3},
    {"label": "A7",  "title": "N = 3 Model Benchmarking Baseline Comparison",        "report": "initial_3",     "table_num": 4},
    {"label": "A8",  "title": "N = 3 Model Benchmarking Runs Justification",         "report": "initial_3",     "table_num": 7},
    {"label": "A9",  "title": "N = 3 Stability Analysis Model Benchmarking",         "report": "temperature_3", "table_num": 1},
    {"label": "A10", "title": "N = 3 Stability Analysis Reliability Metrics",        "report": "temperature_3", "table_num": 2},
    {"label": "A11", "title": "N = 3 Stability Analysis Efficiency Metrics",         "report": "temperature_3", "table_num": 3},
    {"label": "A12", "title": "N = 3 Stability Analysis Baseline Comparison",        "report": "temperature_3", "table_num": 4},
    {"label": "A13", "title": "N = 3 Stability Analysis Variance Analysis",          "report": "temperature_3", "table_num": 7},
    {"label": "A14", "title": "N = 1 Real-World Model Performance Metrics",          "report": "realworld_1",   "table_num": 1},
    {"label": "A15", "title": "N = 1 Real-World Average Failures",                   "report": "realworld_1",   "table_num": 2},
    {"label": "A16", "title": "N = 1 Real-World Efficiency Metrics",                 "report": "realworld_1",   "table_num": 3},
    {"label": "A17", "title": "N = 10 Real-World Model Performance Metrics",         "report": "realworld_10",  "table_num": 1},
    {"label": "A18", "title": "N = 10 Real-World Reliability Metrics",               "report": "realworld_10",  "table_num": 2},
    {"label": "A19", "title": "N = 10 Real-World Efficiency Metrics",                "report": "realworld_10",  "table_num": 3},
    {"label": "A20", "title": "N = 10 Real-World Baseline Comparison",               "report": "realworld_10",  "table_num": 4},
    {"label": "A21", "title": "N = 10 Real-World Pairwise Statistical Comparison",   "report": "realworld_10",  "table_num": 5},
    {"label": "A22", "title": "N = 10 Real-World Bootstrap Ranking Stability",       "report": "realworld_10",  "table_num": 6},
]


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            edge_el.set(qn('w:space'), str(edge_data.get('space', 0)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def parse_table_lines(raw_lines):
    """Convert raw '| cell | cell |' lines into a list of row lists, skipping separator lines."""
    result = []
    for line in raw_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Skip separator rows (only dashes/spaces)
        if not all(set(c.replace('-', '').replace(' ', '')) == set() for c in cells):
            result.append(cells)
    return result


def extract_table(lines, table_num):
    """
    Extract a single table from report lines.

    table_num == 1: no marker — find first '| ' header row (within first 50 lines).
    table_num >= 2: scan for 'TABLE {n}:' marker then capture the following | rows.

    Returns a list-of-lists (rows × cols), or [] if not found.
    """
    if table_num == 1:
        for i, line in enumerate(lines):
            if line.strip().startswith('|') and i < 60:
                raw = []
                j = i
                while j < len(lines):
                    if lines[j].strip().startswith('|'):
                        raw.append(lines[j])
                    elif 'Note:' in lines[j] or lines[j].strip().startswith('='):
                        break
                    j += 1
                return parse_table_lines(raw)
        return []

    marker = f'TABLE {table_num}:'
    for i, line in enumerate(lines):
        if marker in line:
            j = i + 1
            # Skip non-table lines until we hit a '|' row
            while j < len(lines) and not lines[j].strip().startswith('|'):
                j += 1
            if j >= len(lines):
                return []
            raw = []
            while j < len(lines):
                if lines[j].strip().startswith('|'):
                    raw.append(lines[j])
                elif (lines[j].strip().startswith('=') or
                      'Note:' in lines[j] or
                      'FOOTNOTES' in lines[j]) and raw:
                    break
                j += 1
            return parse_table_lines(raw)
    return []


def load_report_lines(path):
    """Read a report file and return its lines, or [] if missing."""
    if not path.exists():
        print(f"  WARNING: {path} not found — skipping tables from this report.")
        return []
    with open(path, 'r', encoding='utf-8') as f:
        return f.readlines()


def create_docx_table(doc, table_data, label, title):
    """
    Add a formatted table to the document.

    - Header row: bold, centered, blue-tinted background.
    - Literature rows (Mode cell contains 'Literature'): model name underlined.
    - Numeric columns (index > 1): center-aligned.
    """
    if not table_data:
        print(f"  WARNING: Table {label} is empty — skipped.")
        return

    # Caption
    caption = doc.add_paragraph()
    caption_run = caption.add_run(f'Table {label}: {title}')
    caption_run.bold = True
    caption_run.font.size = Pt(11)
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

    num_cols = len(table_data[0])
    num_rows = len(table_data)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Column widths
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Inches(2.0) if i == 0 else (Inches(1.5) if i == 1 else Inches(1.0))

    # Locate Mode column index from header
    header = table_data[0]
    mode_col_idx = next((i for i, h in enumerate(header) if h.strip().lower() == "mode"), None)

    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]

        # Determine if this is a literature baseline row
        is_literature = (
            row_idx > 0
            and mode_col_idx is not None
            and mode_col_idx < len(row_data)
            and "Literature" in row_data[mode_col_idx]
        )

        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text

            if row_idx == 0:
                # Header formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9E2F3')
                cell._element.get_or_add_tcPr().append(shading)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        # Underline model name cell for literature rows
                        if is_literature and col_idx == 0:
                            run.font.underline = True
                    if col_idx > 1:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def main():
    base_path = Path(r'C:\Repos\slm-python-unit-test-benchmark\step4_evaluation')

    report_paths = {
        "initial_1":     base_path / 'initial_1_plots'     / 'summary_report.txt',
        "initial_3":     base_path / 'initial_3_plots'     / 'summary_report.txt',
        "temperature_3": base_path / 'temperature_3_plots'  / 'summary_report.txt',
        "realworld_1":   base_path / 'realworld_1_plots'   / 'summary_report.txt',
        "realworld_10":  base_path / 'realworld_10_plots'  / 'summary_report.txt',
    }

    # Load all reports once
    report_lines = {key: load_report_lines(path) for key, path in report_paths.items()}

    doc = Document()

    title_para = doc.add_heading('Evaluation Summary Tables', level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph(
        'Formatted tables extracted from evaluation summary reports. '
        'Literature baseline rows are indicated by underlined model names '
        'and a "Literature (Source)" entry in the Mode column. '
        'Values marked [a] use Execution Correctness (EC) rather than assertion-based Pass@1 '
        '(see footnotes in corresponding summary_report.txt for full definitions).'
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

    created = 0
    for spec in TABLE_SPEC:
        label   = spec["label"]
        title   = spec["title"]
        key     = spec["report"]
        tnum    = spec["table_num"]
        lines   = report_lines.get(key, [])

        print(f"  Table {label}: extracting TABLE {tnum} from {key}...")
        table_data = extract_table(lines, tnum)

        if table_data:
            create_docx_table(doc, table_data, label, title)
            created += 1
            # Page break after every table so each starts on a fresh page
            doc.add_page_break()
        else:
            print(f"    -> Not found / empty, skipped.")

    output_path = base_path / 'summary_tables.docx'
    doc.save(output_path)
    print(f"\n[SUCCESS] {output_path}")
    print(f"  Tables created: {created} / {len(TABLE_SPEC)}")


if __name__ == '__main__':
    main()
