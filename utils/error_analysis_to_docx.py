"""
Convert error_analysis CSV files to formatted DOCX tables for thesis.

This script reads failure analysis CSV files and creates summary tables
showing failure mode distributions and root cause categories by model.
"""

import csv
from collections import defaultdict
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_docx_table(doc, table_data, title, table_number):
    """
    Create a formatted table in the DOCX document.

    Args:
        doc: Document object
        table_data: List of rows (list of lists)
        title: Table title
        table_number: Table number for caption
    """
    if not table_data:
        return

    # Add table caption
    caption = doc.add_paragraph()
    caption_run = caption.add_run(f'Table {table_number}: {title}')
    caption_run.bold = True
    caption_run.font.size = Pt(11)
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Create table
    num_cols = len(table_data[0])
    num_rows = len(table_data)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Set column widths based on content
    if num_cols > 0:
        # Adjust column widths
        for i, cell in enumerate(table.rows[0].cells):
            if i == 0:  # Model name column
                cell.width = Inches(2.0)
            elif i == 1:  # Second column (Config or first data column)
                cell.width = Inches(1.5)
            else:
                cell.width = Inches(1.0)

    # Fill table with data
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)

            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set header background color
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9E2F3')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # Format data rows
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    # Center align numeric columns
                    if col_idx > 1:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after table
    doc.add_paragraph()


def load_csv_data(file_path):
    """Load and parse CSV file."""
    data = []
    with open(file_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            data.append(row)
    return data


def aggregate_by_model_and_failure_mode(data):
    """
    Aggregate failure counts by model and failure mode.

    Returns:
        dict: {model: {failure_mode: count}}
    """
    aggregated = defaultdict(lambda: defaultdict(int))

    for row in data:
        model = row.get('model', 'Unknown')
        failure_mode = row.get('failure_mode', 'Unknown')
        aggregated[model][failure_mode] += 1

    return aggregated


def aggregate_by_model_and_root_cause(data):
    """
    Aggregate failure counts by model and root cause category.

    Returns:
        dict: {model: {root_cause: count}}
    """
    aggregated = defaultdict(lambda: defaultdict(int))

    for row in data:
        model = row.get('model', 'Unknown')
        root_cause = row.get('root_cause_category', 'Unknown')
        if not root_cause or root_cause.strip() == '':
            root_cause = 'Unclassified'
        aggregated[model][root_cause] += 1

    return aggregated


def aggregate_by_model_config_and_failure_mode(data):
    """
    Aggregate failure counts by model+config and failure mode.

    Returns:
        dict: {(model, config): {failure_mode: count}}
    """
    aggregated = defaultdict(lambda: defaultdict(int))

    for row in data:
        model = row.get('model', 'Unknown')
        config = row.get('config', 'Unknown')
        failure_mode = row.get('failure_mode', 'Unknown')
        aggregated[(model, config)][failure_mode] += 1

    return aggregated


def aggregate_by_model_config_and_root_cause(data):
    """
    Aggregate failure counts by model+config and root cause category.

    Returns:
        dict: {(model, config): {root_cause: count}}
    """
    aggregated = defaultdict(lambda: defaultdict(int))

    for row in data:
        model = row.get('model', 'Unknown')
        config = row.get('config', 'Unknown')
        root_cause = row.get('root_cause_category', 'Unknown')
        if not root_cause or root_cause.strip() == '':
            root_cause = 'Unclassified'
        aggregated[(model, config)][root_cause] += 1

    return aggregated


def aggregate_by_proximity(data):
    """
    Aggregate failure counts by model and proximity to success.

    Returns:
        dict: {model: {proximity: count}}
    """
    aggregated = defaultdict(lambda: defaultdict(int))

    for row in data:
        model = row.get('model', 'Unknown')
        proximity = row.get('proximity_to_success', 'Unknown')
        if not proximity or proximity.strip() == '':
            proximity = 'Unclassified'
        aggregated[model][proximity] += 1

    return aggregated


def create_failure_mode_table(aggregated_data):
    """Create table data for failure mode distribution."""
    # Get all unique failure modes
    all_modes = set()
    for model_data in aggregated_data.values():
        all_modes.update(model_data.keys())
    all_modes = sorted(all_modes)

    # Create header row
    header = ['Model'] + all_modes + ['Total']

    # Create data rows
    rows = [header]
    for model in sorted(aggregated_data.keys()):
        row = [model]
        total = 0
        for mode in all_modes:
            count = aggregated_data[model].get(mode, 0)
            row.append(count)
            total += count
        row.append(total)
        rows.append(row)

    return rows


def create_failure_mode_table_with_config(aggregated_data):
    """Create table data for failure mode distribution with config."""
    # Get all unique failure modes
    all_modes = set()
    for model_data in aggregated_data.values():
        all_modes.update(model_data.keys())
    all_modes = sorted(all_modes)

    # Create header row
    header = ['Model', 'Config'] + all_modes + ['Total']

    # Create data rows
    rows = [header]
    for (model, config) in sorted(aggregated_data.keys()):
        row = [model, config]
        total = 0
        for mode in all_modes:
            count = aggregated_data[(model, config)].get(mode, 0)
            row.append(count)
            total += count
        row.append(total)
        rows.append(row)

    return rows


def create_root_cause_table(aggregated_data):
    """Create table data for root cause distribution."""
    # Get all unique root causes
    all_causes = set()
    for model_data in aggregated_data.values():
        all_causes.update(model_data.keys())
    all_causes = sorted(all_causes)

    # Create header row
    header = ['Model'] + all_causes + ['Total']

    # Create data rows
    rows = [header]
    for model in sorted(aggregated_data.keys()):
        row = [model]
        total = 0
        for cause in all_causes:
            count = aggregated_data[model].get(cause, 0)
            row.append(count)
            total += count
        row.append(total)
        rows.append(row)

    return rows


def create_root_cause_table_with_config(aggregated_data):
    """Create table data for root cause distribution with config."""
    # Get all unique root causes
    all_causes = set()
    for model_data in aggregated_data.values():
        all_causes.update(model_data.keys())
    all_causes = sorted(all_causes)

    # Create header row
    header = ['Model', 'Config'] + all_causes + ['Total']

    # Create data rows
    rows = [header]
    for (model, config) in sorted(aggregated_data.keys()):
        row = [model, config]
        total = 0
        for cause in all_causes:
            count = aggregated_data[(model, config)].get(cause, 0)
            row.append(count)
            total += count
        row.append(total)
        rows.append(row)

    return rows


def create_proximity_table(aggregated_data):
    """Create table data for proximity to success distribution."""
    # Get all unique proximity values
    all_proximities = set()
    for model_data in aggregated_data.values():
        all_proximities.update(model_data.keys())
    all_proximities = sorted(all_proximities)

    # Create header row
    header = ['Model'] + all_proximities + ['Total']

    # Create data rows
    rows = [header]
    for model in sorted(aggregated_data.keys()):
        row = [model]
        total = 0
        for proximity in all_proximities:
            count = aggregated_data[model].get(proximity, 0)
            row.append(count)
            total += count
        row.append(total)
        rows.append(row)

    return rows


def process_error_analysis_file(file_path, doc, report_name, start_table_num):
    """
    Process a single error analysis CSV file.

    Args:
        file_path: Path to CSV file
        doc: Document object
        report_name: Name of the report section
        start_table_num: Starting table number

    Returns:
        Next available table number
    """
    data = load_csv_data(file_path)

    if not data:
        print(f"  Warning: No data found in {file_path}")
        return start_table_num

    # Add section heading
    doc.add_heading(report_name, level=1)

    # Add summary info
    summary = doc.add_paragraph()
    summary.add_run(f"Total failures analyzed: {len(data)}\n")
    summary.add_run('\n')

    table_num = start_table_num

    # Table 1: Failure Mode Distribution by Model
    failure_mode_agg = aggregate_by_model_and_failure_mode(data)
    failure_mode_table = create_failure_mode_table(failure_mode_agg)
    if len(failure_mode_table) > 1:
        create_docx_table(
            doc,
            failure_mode_table,
            'Failure Mode Distribution by Model',
            table_num
        )
        table_num += 1

    # Table 2: Failure Mode Distribution by Model and Config
    failure_mode_config_agg = aggregate_by_model_config_and_failure_mode(data)
    failure_mode_config_table = create_failure_mode_table_with_config(failure_mode_config_agg)
    if len(failure_mode_config_table) > 1:
        create_docx_table(
            doc,
            failure_mode_config_table,
            'Failure Mode Distribution by Model and Configuration',
            table_num
        )
        table_num += 1

    # Table 3: Root Cause Distribution by Model
    root_cause_agg = aggregate_by_model_and_root_cause(data)
    root_cause_table = create_root_cause_table(root_cause_agg)
    if len(root_cause_table) > 1:
        create_docx_table(
            doc,
            root_cause_table,
            'Root Cause Category Distribution by Model',
            table_num
        )
        table_num += 1

    # Table 4: Root Cause Distribution by Model and Config
    root_cause_config_agg = aggregate_by_model_config_and_root_cause(data)
    root_cause_config_table = create_root_cause_table_with_config(root_cause_config_agg)
    if len(root_cause_config_table) > 1:
        create_docx_table(
            doc,
            root_cause_config_table,
            'Root Cause Category Distribution by Model and Configuration',
            table_num
        )
        table_num += 1

    # Table 5: Proximity to Success Distribution
    proximity_agg = aggregate_by_proximity(data)
    proximity_table = create_proximity_table(proximity_agg)
    if len(proximity_table) > 1:
        create_docx_table(
            doc,
            proximity_table,
            'Proximity to Success Distribution by Model',
            table_num
        )
        table_num += 1

    # Add page break after each report
    doc.add_page_break()

    return table_num


def main():
    """Main function to process all error analysis files."""
    # Define paths
    base_path = Path(r'C:\Repos\slm-python-unit-test-benchmark\step4_evaluation')
    error_analysis_path = base_path / 'error_analysis_output'

    reports = [
        {
            'path': error_analysis_path / 'realworld_failure_analysis.csv',
            'name': 'Real World Failure Analysis'
        },
        {
            'path': error_analysis_path / 'testeval_initial_failure_analysis.csv',
            'name': 'TestEval Initial Failure Analysis'
        },
        {
            'path': error_analysis_path / 'testeval_temperature_failure_analysis.csv',
            'name': 'TestEval Temperature Failure Analysis'
        }
    ]

    # Create document
    doc = Document()

    # Set document title
    title = doc.add_heading('Error Analysis Summary Tables', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add document description
    intro = doc.add_paragraph()
    intro.add_run(
        'This document contains formatted tables summarizing failure analysis data. '
        'These tables present failure mode distributions, root cause categories, and '
        'proximity to success metrics for various language models evaluated on unit test '
        'generation benchmarks.'
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Process each report
    table_counter = 1
    for report in reports:
        if report['path'].exists():
            print(f"Processing {report['name']}...")
            table_counter = process_error_analysis_file(
                report['path'],
                doc,
                report['name'],
                table_counter
            )
        else:
            print(f"Warning: {report['path']} not found, skipping...")

    # Save document
    output_path = error_analysis_path / 'error_analysis_tables.docx'
    doc.save(output_path)

    print(f"\n[SUCCESS] Document created successfully: {output_path}")
    print(f"  Total tables created: {table_counter - 1}")
    print("\nYou can now open the document and copy the tables to your thesis.")


if __name__ == '__main__':
    main()
